import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def add_code_to_document(doc, file_path, code):
    # Replace "sap" with "fba" in the file path for the header
    file_path = file_path.replace("sap", "fba")

    # Replace "sap" with "fba" only in lines starting with "package" or "import"
    code_lines = code.splitlines()
    updated_code_lines = []
    for line in code_lines:
        if line.startswith("package") or line.startswith("import"):
            updated_code_lines.append(line.replace("sap", "fba"))
        else:
            updated_code_lines.append(line)
    updated_code = "\n".join(updated_code_lines)

    # Add the file path as a heading with font size 12
    heading = doc.add_heading(level=2)
    run = heading.add_run(file_path)
    run.font.size = Pt(12)

    # Add the code with font size 12 and line spacing 1.15
    p = doc.add_paragraph()
    run = p.add_run(updated_code)
    run.font.size = Pt(12)
    p.line_spacing_rule = 1.15
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT


def set_document_margins(doc, top=1, bottom=1, left=1, right=1):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def read_code_files(directory):
    code_files = []
    for root, _, files in os.walk(directory):
        for file in files:
            if file.endswith('.java'):  # you can add other extensions if needed
                file_path = os.path.join(root, file)
                with open(file_path, 'r', encoding='utf-8') as f:
                    code = f.read()
                relative_path = os.path.relpath(file_path, directory)
                code_files.append((relative_path, code))
    return code_files


def create_word_documents(code_files, output_prefix):
    output_dir = 'output/autoconfigure'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document()

    # Set page margins to 1 inch
    set_document_margins(doc)

    current_file_index = 1
    current_page_count = 0

    for file_path, code in code_files:
        add_code_to_document(doc, file_path, code)
        current_page_count += doc.element.xpath('count(//w:sectPr)')

        if current_page_count >= 24:
            output_file = os.path.join(output_dir, f"{output_prefix}_{current_file_index}.docx")
            doc.save(output_file)
            print(f"Saved {output_file}")
            doc = Document()  # Start a new document
            set_document_margins(doc)
            current_file_index += 1
            current_page_count = 0

    # Save remaining pages, if any
    if current_page_count > 0:
        output_file = os.path.join(output_dir, f"{output_prefix}_{current_file_index}.docx")
        doc.save(output_file)
        print(f"Saved {output_file}")


def main():
    input_directory = 'input/app-frw-ng-plus/sub-projects/app-frw-autoconfigure'  # replace with your path
    output_prefix = 'app_code'
    code_files = read_code_files(input_directory)
    create_word_documents(code_files, output_prefix)


if __name__ == '__main__':
    main()
